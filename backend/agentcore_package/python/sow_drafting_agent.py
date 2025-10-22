# sow_drafting_agent_claude.py
import boto3, json, os, tempfile, uuid, concurrent.futures
from datetime import datetime
from docx import Document
from strands import Agent

class SOWDraftingAgent(Agent):
    def __init__(self, name="sow-drafting-agent", region="us-east-1"):
        super().__init__(name=name)
        self.region = region
        self.s3 = boto3.client("s3", region_name=region)
        self.bedrock = boto3.client("bedrock-runtime", region_name=region)
        self.profile_arn = self._find_inference_profile("anthropic.claude-3-5-sonnet-20241022-v2:0")
        print(f"[INFO] ✅ Using Claude 3.5 Sonnet profile: {self.profile_arn}")

    # ---------- Utility ----------
    def _find_inference_profile(self, model_id):
        bedrock = boto3.client("bedrock", region_name=self.region)
        try:
            profiles = bedrock.list_inference_profiles()["inferenceProfileSummaries"]
            for p in profiles:
                if model_id in p.get("modelArn", "") or model_id in p.get("inferenceProfileArn", ""):
                    return p["inferenceProfileArn"]
        except Exception as e:
            print(f"[WARN] Could not list inference profiles: {e}")
        raise Exception(f"No inference profile found for {model_id}")

    def read_json_from_s3(self, bucket, key):
        obj = self.s3.get_object(Bucket=bucket, Key=key)
        return json.loads(obj["Body"].read().decode("utf-8"))

    def save_docx_to_s3(self, document, bucket, key):
        tmp = tempfile.mktemp(suffix=".docx")
        document.save(tmp)
        self.s3.upload_file(tmp, bucket, key)
        os.remove(tmp)
        return f"s3://{bucket}/{key}"

    # ---------- Claude + Titan Section Generator ----------
    def generate_section(self, title, context):
        prompt = f"""
        You are a **Senior Presales Consultant** drafting a customer-facing Statement of Work (SOW).

        Task:
        Write the **{title}** section of the SOW using the following project context.

        Context (RFP + Clarifications + Pricing):
        {json.dumps(context, indent=2)[:8000]}

        Guidelines:
        - Be formal, clear, and customer-centric.
        - Structure with 1–3 concise paragraphs.
        - Avoid filler, disclaimers, or restating the title.
        - Use confident tone and business clarity.
        Return only the text (no headings or markdown).
        """

        # --- Try Claude 3.5 Sonnet ---
        try:
            payload = {
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 1200,
                "temperature": 0.4,
                "messages": [{"role": "user", "content": prompt}],
            }

            resp = self.bedrock.invoke_model(
                modelId=self.profile_arn,
                body=json.dumps(payload),
                accept="application/json",
                contentType="application/json",
            )

            raw = resp["body"].read().decode("utf-8")
            out = json.loads(raw)
            if "content" in out:
                return out["content"][0]["text"].strip()
        except Exception as e:
            print(f"[WARN] Claude 3.5 failed for '{title}': {e}")

        # --- Fallback: Titan Express ---
        try:
            payload = {
                "inputText": prompt,
                "textGenerationConfig": {"temperature": 0.3, "maxTokenCount": 1200},
            }
            resp = self.bedrock.invoke_model(
                modelId="amazon.titan-text-express-v1",
                body=json.dumps(payload),
                accept="application/json",
                contentType="application/json",
            )
            out = json.loads(resp["body"].read().decode("utf-8"))
            if "results" in out:
                return out["results"][0]["outputText"].strip()
        except Exception as e:
            print(f"[WARN] Titan fallback failed for '{title}': {e}")

        return f"[Placeholder] Unable to generate section: {title}"

    # ---------- Main Agent Logic ----------
    def run(self, bucket, parsed_key, clar_key, pricing_key, bucket_out):
        print(f"[INFO] Running SOWDraftingAgent (Claude 3.5) for user data under: {parsed_key}")

        # Step 1: Load all input data
        parsed_data = self.read_json_from_s3(bucket, parsed_key)
        clar_data = self.read_json_from_s3(bucket, clar_key)
        pricing_data = self.read_json_from_s3(bucket, pricing_key)

        customer = parsed_data.get("customer_name", "Client Organization")
        project_title = parsed_data.get("project_title", "Proposed Project")

        context = {
            "parsed_rfp": parsed_data,
            "clarifications": clar_data,
            "pricing_summary": pricing_data,
        }

        print("[INFO] Generating SOW sections via Claude 3.5 Sonnet...")

        section_titles = [
            "Project Overview",
            "Objectives",
            "Scope of Work",
            "Deliverables",
            "Timeline & Milestones",
            "Pricing & Payment Terms",
            "Assumptions",
            "Roles & Responsibilities",
            "Acceptance Criteria",
            "Terms & Conditions",
        ]

        # Step 2: Generate sections concurrently
        sections = {}
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            future_to_title = {executor.submit(self.generate_section, t, context): t for t in section_titles}
            for future in concurrent.futures.as_completed(future_to_title):
                title = future_to_title[future]
                try:
                    sections[title] = future.result()
                except Exception as e:
                    sections[title] = f"[Error generating section: {title}] ({e})"

        # Step 3: Create the DOCX file
        print("[INFO] Creating SOW DOCX file...")
        doc = Document()
        doc.add_heading("Statement of Work (SOW)", level=1)
        doc.add_paragraph(f"Client: {customer}")
        doc.add_paragraph(f"Project: {project_title}")
        doc.add_paragraph(f"Generated on: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph("\n")

        for title, text in sections.items():
            doc.add_heading(title, level=2)
            doc.add_paragraph(text)
            doc.add_paragraph("\n")

        # Step 4: Save to S3
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        user_prefix = parsed_key.split("/")[0]
        out_folder = f"{user_prefix}/sow_drafts/"
        safe_project = project_title.replace(" ", "_").replace("/", "_")
        out_key = f"{out_folder}{safe_project}_SOW_{ts}.docx"

        s3_uri = self.save_docx_to_s3(doc, bucket_out, out_key)
        print(f"[INFO] ✅ SOW Draft saved to {s3_uri}")
        return out_key


# ---------- Local Test ----------
if __name__ == "__main__":
    agent = SOWDraftingAgent(region="us-east-1")
    out = agent.run(
        bucket="presales-rfp-outputs",
        parsed_key="ravi/parsed_outputs/RFP_1_20251013_045445_parsed.json",
        clar_key="ravi/clarifications/RFP_1_20251013_045445_parsed_clarifications_20251013_045454.json",
        pricing_key="ravi/pricing_outputs/RFP_1_20251013_045445_parsed_pricing_20251013_190854.json",
        bucket_out="presales-rfp-outputs"
    )
    print("✅ Output:", out)
